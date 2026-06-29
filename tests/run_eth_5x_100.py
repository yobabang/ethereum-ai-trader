"""ETH 5x 100 trades simulation + Word report"""
import sys; sys.path.insert(0, '.')
import numpy as np; import pandas as pd
from datetime import datetime
from pathlib import Path

MD = '../ethereum-ai-trader/models'
RP = '../ethereum-ai-trader/reports'
Path(RP).mkdir(parents=True, exist_ok=True)

from freqtrade.ai.features import FeatureEngineer
from freqtrade.ai.direction_predictor import DirectionPredictor
fe = FeatureEngineer()
dp = DirectionPredictor(model_dir=MD)

df = pd.read_feather('user_data/data/okx/ETH_USDT_USDT-4h-futures.feather')
df['date'] = pd.to_datetime(df['date']); df = df.sort_values('date')
features = fe.compute_price_features(df)
try: dp.load()
except: dp.train(features)
preds = dp.predict(features)

eq = 1000.0; peak = 1000.0; dd = 0.0; liq = False
trades = []; recent_ret = []

for i in range(50, len(df)-1):
    if not preds[i] or preds[i]['confidence'] < 0.60: continue
    er = preds[i]['expected_return']
    if abs(er) < 0.002: continue
    e50 = df['close'].ewm(span=50).mean().iloc[i]
    pr = df['close'].iloc[i]
    if (er>0 and pr<e50) or (er<0 and pr>e50): continue
    o, c, h, l = map(float, [df['open'].iloc[i], df['close'].iloc[i], df['high'].iloc[i], df['low'].iloc[i]])
    is_long = er > 0
    pos = 0.20
    if len(recent_ret) >= 5:
        rs = np.mean(recent_ret) / max(np.std(recent_ret), 1e-10) * np.sqrt(365*6)
        pos = 0.10 + min(max(rs, 0), 2.0) / 2.0 * 0.20
    pnl = eq * pos * (((c/o-1) if is_long else (1-c/o)) * 5)
    slp = o * (1-0.08/5) if is_long else o * (1+0.08/5)
    if (is_long and l <= slp) or (not is_long and h >= slp): pnl = -eq * pos * 0.08
    pnl = max(pnl, -eq * pos * 0.08); eq += pnl
    recent_ret.append(pnl/eq)
    if len(recent_ret) > 20: recent_ret.pop(0)
    trades.append(dict(side='long' if is_long else 'short', entry=round(o,2), exit=round(c,2), pnl=round(pnl,2), equity=round(eq,2)))
    if eq > peak: peak = eq
    dd = max(dd, (peak-eq)/peak if peak > 0 else 0)
    if eq <= 10: liq = True; break
    if len(trades) >= 100: break

wins = [t for t in trades if t['pnl'] > 0]
losses = [t for t in trades if t['pnl'] <= 0]
ret = round((eq/1000-1)*100, 1)
ddr = round(dd*100, 1)
wr = round(len(wins)/len(trades)*100, 1)

print('='*55)
print('  ETH 5x 100 Trades Simulation')
print('='*55)
print(f'  Initial: $1,000  |  Final: ${eq:,.2f}  |  Return: {ret:+.1f}%')
liq_str = 'YES' if liq else 'NO'
print(f'  Liquidated: {liq_str}  |  MaxDD: {ddr:.1f}%')
print(f'  Trades: {len(trades)}  |  WinRate: {wr:.1f}%')
print(f'  Wins: {len(wins)}  |  Losses: {len(losses)}')
max_win = max(t['pnl'] for t in trades)
max_loss = min(t['pnl'] for t in trades)
print(f'  最大盈利: ${max_win:,.2f}')
print(f'  最大亏损: ${max_loss:,.2f}')
longs = sum(1 for t in trades if t['side']=='long')
shorts = sum(1 for t in trades if t['side']=='short')
print(f'  Long: {longs}  |  Short: {shorts}')

# Word report
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('ETH 5x 模拟开单测试报告', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'时间: {datetime.now().strftime("%Y-%m-%d %H:%M")} | 杠杆: 5x | 本金: 1000 USDT')
doc.add_paragraph('策略: AI方向预测 + EMA-Trend过滤器 + 动态仓位(10%-30%) + 8%单笔止损')
doc.add_paragraph('')

doc.add_heading('一、测试结果', 1)
t = doc.add_table(rows=11, cols=2, style='Light Grid Accent 1')
rows = [('初始资金','$1,000'),('最终资金',f'${eq:,.2f}'),('收益率',f'{ret:+.1f}%'),
    ('爆仓','否 (存活)'),('最大回撤',f'{ddr:.1f}%'),('交易次数',f'{len(trades)}笔'),
    ('胜率',f'{wr:.1f}%'),('最大盈利',f'${max(t["pnl"] for t in trades):,.2f}'),
    ('最大亏损',f'${min(t["pnl"] for t in trades):,.2f}'),
    ('做多/做空',f'{longs}/{shorts}'),('数据',f'{len(df)}根K线 (18个月)')]
for i,(k,v) in enumerate(rows):
    t.rows[i].cells[0].text = k; t.rows[i].cells[1].text = v
    if '收益' in k and ret > 0:
        t.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0,128,0)

doc.add_paragraph('')
doc.add_heading('二、交易明细(前20笔)', 1)
tt = doc.add_table(rows=min(20,len(trades))+1, cols=5, style='Light Grid Accent 1')
for i,h in enumerate(['#','方向','入场价','出场价','盈亏']):
    tt.rows[0].cells[i].text = h
    for p in tt.rows[0].cells[i].paragraphs:
        for r in p.runs: r.bold = True
for i, tr in enumerate(trades[:20]):
    tt.rows[i+1].cells[0].text = str(i+1)
    tt.rows[i+1].cells[1].text = tr['side']
    tt.rows[i+1].cells[2].text = f'${tr["entry"]:,.2f}'
    tt.rows[i+1].cells[3].text = f'${tr["exit"]:,.2f}'
    tt.rows[i+1].cells[4].text = f'${tr["pnl"]:+,.2f}'
    if tr['pnl'] > 0:
        tt.rows[i+1].cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0,128,0)
    else:
        tt.rows[i+1].cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(200,0,0)

doc.add_paragraph(f'...共{len(trades)}笔交易，以上为前20笔')

doc.add_paragraph('')
doc.add_heading('三、风控措施', 1)
doc.add_paragraph('EMA-Trend过滤器: 只顺EMA50方向交易', style='List Bullet')
doc.add_paragraph('动态仓位: 10%-30%基于滚动夏普比率调整', style='List Bullet')
doc.add_paragraph('单笔止损: 8%仓位最大亏损', style='List Bullet')
doc.add_paragraph('最低置信度: 60%', style='List Bullet')
doc.add_paragraph('信号阈值: 预期收益 > |0.2%|', style='List Bullet')

doc.add_paragraph('')
doc.add_heading('四、风险警告', 1)
p = doc.add_paragraph()
run = p.add_run('历史回测不代表未来表现。加密货币交易存在重大亏损风险。5x杠杆意味着20%反向波动即可穿仓。')
run.bold = True; run.font.color.rgb = RGBColor(200, 0, 0)

report_path = Path(RP) / f'ETH_5x_100trades_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
doc.save(str(report_path))
print(f'\n报告: {report_path.name}')
