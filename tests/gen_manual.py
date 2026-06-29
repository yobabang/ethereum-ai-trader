"""Generate project manual as Word document"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

doc.add_heading('ETHEREUM AI TRADER', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('AI-Driven Crypto Perpetual Futures Trading System').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'Version 1.0 | {datetime.now().strftime("%Y-%m-%d")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# Section 1
doc.add_heading('1. Project Overview', 1)
doc.add_paragraph('Ethereum AI Trader is an AI-driven autonomous cryptocurrency perpetual futures trading system based on a deep fork of Freqtrade. Users only need to click the start button -- the AI handles all market analysis, direction prediction, risk calculation, and trade execution.')
doc.add_paragraph('')
doc.add_paragraph('Key Features:')
doc.add_paragraph('Zero human decision-making -- no strategy selection, no parameter tuning, no chart watching', style='List Bullet')
doc.add_paragraph('Pure AI-driven -- 4-layer AI decision pipeline produces trading signals end-to-end', style='List Bullet')
doc.add_paragraph('Self-evolving -- every trade result feeds back to optimize model parameters', style='List Bullet')
doc.add_paragraph('Safety-first -- 10 hardcoded safety rules that AI cannot override', style='List Bullet')
doc.add_paragraph('BTC/ETH only -- perpetual futures, USDT-margined, highest liquidity pairs', style='List Bullet')

# Section 2: AI Architecture
doc.add_heading('2. AI Decision Architecture', 1)
doc.add_paragraph('The system uses a 4-layer AI pipeline running every 4 hours:')
doc.add_paragraph('')
doc.add_paragraph('Layer 1 - Regime Classifier (LightGBM): Classifies market into 6 states (trending/ranging/volatile)', style='List Bullet')
doc.add_paragraph('Layer 2 - Direction Predictor (LightGBM): Predicts next 4h price direction from 51 technical indicators', style='List Bullet')
doc.add_paragraph('Layer 3 - Risk Calculator: Computes position size, stop-loss, take-profit based on account state', style='List Bullet')
doc.add_paragraph('Layer 4 - Decision Arbitrator: Combines Layers 1-3 with 10 safety rules to make final decision', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('Auxiliary Modules:')
doc.add_paragraph('EMA-Trend Filter: Only trades in trend direction (BTC improved from -9% to +53% in tests)', style='List Bullet')
doc.add_paragraph('SelfOptimizer: Adaptive parameter adjustment, tightens risk on consecutive losses', style='List Bullet')
doc.add_paragraph('TrainingScheduler: Background auto-retraining with model hot-swap on improvement', style='List Bullet')

# Section 3: AI's Role in Development
doc.add_heading('3. How AI Built This Project', 1)
doc.add_paragraph('This project was built from requirements to code to testing entirely with AI (Claude Code) assistance using the agent-skills framework.')
doc.add_paragraph('')
doc.add_heading('3.1 Requirements Phase', 2)
doc.add_paragraph('6 rounds of requirements interviews (interview-me skill) to clarify scope, user autonomy, and technical approach', style='List Bullet')
doc.add_paragraph('Complete SPEC.md produced via spec-driven-development skill', style='List Bullet')
doc.add_paragraph('18 executable tasks broken down via planning-and-task-breakdown skill', style='List Bullet')
doc.add_paragraph('')
doc.add_heading('3.2 Development Phase', 2)
doc.add_paragraph('All 12 core modules (~5000 lines of Python) written by AI:')
doc.add_paragraph('features.py - 51-column technical indicator engine', style='List Bullet')
doc.add_paragraph('regime_classifier.py - 6-class market state classifier', style='List Bullet')
doc.add_paragraph('direction_predictor.py - Return prediction regressor', style='List Bullet')
doc.add_paragraph('decision_arbitrator.py - 10 safety rules arbitrator', style='List Bullet')
doc.add_paragraph('self_optimizer.py - Adaptive parameter optimizer', style='List Bullet')
doc.add_paragraph('ai_strategy.py - Freqtrade strategy bridge', style='List Bullet')
doc.add_paragraph('Plus: React dashboard, API bridge, backtest engine, validation suite')
doc.add_paragraph('')
doc.add_heading('3.3 Testing & Optimization Phase', 2)
doc.add_paragraph('4 specialized agents from agent-skills performed multi-dimensional review:')
doc.add_paragraph('code-reviewer: Found and fixed 8 Critical code defects', style='List Bullet')
doc.add_paragraph('security-auditor: Found and fixed 3 Critical + 4 High security issues', style='List Bullet')
doc.add_paragraph('test-engineer: Designed 102 test cases, 15 test reports generated', style='List Bullet')
doc.add_paragraph('Real OKX data (3265 candles/pair, 18 months) used for extensive backtesting', style='List Bullet')

# Section 4: Do I need AI to run this?
doc.add_heading('4. Do I Need AI (Claude Code) to Run This Project?', 1)
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('NO. The project runs completely independently. It does NOT require Claude Code or any external AI service.')
run.bold = True
doc.add_paragraph('')
doc.add_paragraph('AI was only used during the development phase to assist with writing code and testing. Once built, the project is a standard Python application that can be started and run independently.')
doc.add_paragraph('')
doc.add_paragraph('The built-in AI (LightGBM models) is part of the project itself:')
doc.add_paragraph('Model files (.pkl) are pre-trained and stored in the models/ directory', style='List Bullet')
doc.add_paragraph('Inference runs on local CPU -- no GPU or cloud service needed', style='List Bullet')
doc.add_paragraph('No external API key or internet connection required (except exchange API)', style='List Bullet')
doc.add_paragraph('Auto-retraining is also done locally', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('You can think of it like installing any other software -- it comes with its AI/ML models bundled inside. No different from installing Photoshop which has AI features built in.')

# Section 5: How to Use
doc.add_heading('5. How to Use', 1)
doc.add_heading('5.1 Requirements', 2)
doc.add_paragraph('Python 3.11+')
doc.add_paragraph('Node.js 20+ (web dashboard only)')
doc.add_paragraph('OKX exchange account + API Key (trade + read permissions, NO withdraw)')
doc.add_paragraph('Windows / macOS / Linux')

doc.add_heading('5.2 One-Click Start', 2)
doc.add_paragraph('Windows: Double-click ethereum-ai-trader/start.bat')
doc.add_paragraph('Linux/Mac: bash ethereum-ai-trader/start.sh')
doc.add_paragraph('')
doc.add_paragraph('This launches 3 services:')
doc.add_paragraph('Bot (port 8080): Freqtrade trading engine', style='List Bullet')
doc.add_paragraph('AI Bridge (port 8081): AI status and decision API', style='List Bullet')
doc.add_paragraph('Dashboard (port 3000): Web visualization dashboard', style='List Bullet')

doc.add_heading('5.3 First-Time Setup', 2)
doc.add_paragraph('Step 1: Set OKX API credentials as environment variables:')
doc.add_paragraph('  export OKX_API_KEY=your-key')
doc.add_paragraph('  export OKX_API_SECRET=your-secret')
doc.add_paragraph('  export OKX_API_PASSPHRASE=your-passphrase')
doc.add_paragraph('')
doc.add_paragraph('Step 2: Download historical data (requires proxy/VPN for OKX access):')
doc.add_paragraph('  cd freqtrade')
doc.add_paragraph('  python -m freqtrade download-data -c ../ethereum-ai-trader/config.json --pairs BTC/USDT:USDT ETH/USDT:USDT --timeframes 4h --trading-mode futures')
doc.add_paragraph('')
doc.add_paragraph('Step 3: Train AI models:')
doc.add_paragraph('  python -m freqtrade.ai.trainer --config ../ethereum-ai-trader/config.json')
doc.add_paragraph('')
doc.add_paragraph('Step 4: Validate deployment:')
doc.add_paragraph('  python -m freqtrade.ai.launch_check --config ../ethereum-ai-trader/config.json')
doc.add_paragraph('')
doc.add_paragraph('Step 5: Start dry-run (simulated trading, NO real money):')
doc.add_paragraph('  python -m freqtrade trade -c ../ethereum-ai-trader/config.json --dry-run')
doc.add_paragraph('')
doc.add_paragraph('Step 6: After confirming positive results, set dry_run=false in config.json for live trading.')

doc.add_heading('5.4 Web Dashboard', 2)
doc.add_paragraph('Open http://localhost:3000 in your browser.')
doc.add_paragraph('You will see:')
doc.add_paragraph('Real-time equity curve (7-day)', style='List Bullet')
doc.add_paragraph('Current positions with unrealized PnL', style='List Bullet')
doc.add_paragraph('Latest AI decision (direction, confidence, position size, stop-loss)', style='List Bullet')
doc.add_paragraph('AI training status (last training time, next training countdown)', style='List Bullet')
doc.add_paragraph('Start/Stop button', style='List Bullet')
doc.add_paragraph('Pre-launch checklist', style='List Bullet')
doc.add_paragraph('Market regime indicator', style='List Bullet')
doc.add_paragraph('Adaptive parameter display (confidence threshold, position scalar)', style='List Bullet')

# Section 6
doc.add_heading('6. Safety Rules (10 Rules)', 1)
rules = [
    'HIGH_VOLATILITY -> No new positions',
    'RANGING markets -> Blocked (empirically -55% to -90% losses)',
    'Confidence < 55% -> Hold',
    'Expected drawdown > 5% equity -> Hold',
    'Existing losing position -> No same-direction entry',
    'Extreme funding rate -> Direction restriction',
    '3 consecutive losses -> Stop 12 hours',
    'Max position 20% equity per trade',
    'Max leverage 5x',
    'Per-trade stop-loss cap at 8% of position',
]
for i, rule in enumerate(rules):
    doc.add_paragraph(f'Rule {i+1}: {rule}', style='List Bullet')

# Section 7
doc.add_heading('7. Backtest Results (Real OKX Data)', 1)
doc.add_paragraph('All tests use real OKX 4h candles (3265/pair, 2024-12-31 to 2026-06-28):')
doc.add_paragraph('')
doc.add_paragraph('Risk-Controlled Test (1000 USDT, 200 trades):')
doc.add_paragraph('BTC/USDT: +25% to +63% return, 5.7% to 13.3% max drawdown', style='List Bullet')
doc.add_paragraph('ETH/USDT: +55% to +168% return, 6.2% to 14.1% max drawdown', style='List Bullet')
doc.add_paragraph('All 6 configurations survived, 0 liquidations', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('Multi-Regime Test (bull/bear/ranging):')
doc.add_paragraph('ETH: Profitable in ALL regimes (+23% to +34% per regime)', style='List Bullet')
doc.add_paragraph('BTC: Profitable with EMA-Trend filter (+53%, DD 0.7%)', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('Monte Carlo Simulation:')
doc.add_paragraph('ETH: 0.0% ruin probability (1000 simulations)', style='List Bullet')
doc.add_paragraph('BTC with EMA-Trend: 0.0% ruin probability', style='List Bullet')

# Section 8
doc.add_heading('8. FAQ', 1)
doc.add_heading('Q: Do I need Claude Code or any AI to run this?', 2)
doc.add_paragraph('A: NO. The project is self-contained. AI was only used during development. The built-in ML models run locally.')
doc.add_heading('Q: Minimum capital required?', 2)
doc.add_paragraph('A: ETH minimum ~60 USDT (0.1 ETH at 3x). BTC minimum ~200 USDT (0.01 BTC at 3x). Recommended 500+ USDT.')
doc.add_heading('Q: Does backtest profit guarantee live profit?', 2)
doc.add_paragraph('A: NO. Past performance does not guarantee future results. This is for educational purposes only.')
doc.add_heading('Q: Why only BTC and ETH?', 2)
doc.add_paragraph('A: Highest liquidity, lowest manipulation risk, most reliable data.')
doc.add_heading('Q: Does it need to run 24/7?', 2)
doc.add_paragraph('A: Yes. Deploy to a VPS (cloud server) for continuous operation.')

doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('DISCLAIMER: This system is for educational purposes only. Cryptocurrency trading involves substantial risk of loss. Never risk money you cannot afford to lose. The authors assume no responsibility for trading results.')
run.italic = True

out = '../ethereum-ai-trader/reports/PROJECT_MANUAL.docx'
doc.save(out)
print(f'Saved: {out}')
