"""Convert all JSON iteration reports to Word documents."""
import json, glob
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORTS_DIR = Path('../ethereum-ai-trader/reports')

def convert_json_to_word(json_path, docx_path):
    with open(json_path) as f:
        data = json.load(f)

    doc = Document()
    doc.add_heading('以太 AI Trader — 测试报告', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    ts = data.get('timestamp', 'N/A')
    doc.add_paragraph(f'时间: {ts} | 数据: OKX 永续合约 4h | 本金: 1000 USDT | 200笔/配置')
    doc.add_paragraph('')

    # Results table
    results = data.get('results', [])
    if isinstance(results, dict):
        results = [v for v in results.values() if isinstance(v, dict)]
    if not isinstance(results, list):
        results = []
    if results and isinstance(results[0], dict):
        doc.add_heading('一、测试结果', 1)
        table = doc.add_table(rows=len(results)+1, cols=8, style='Light Grid Accent 1')
        headers = ['交易对', '杠杆', '仓位', '止损', '最终资金', '收益率', '最大回撤', '胜率']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            for p in table.rows[0].cells[i].paragraphs:
                for r in p.runs: r.bold = True

        for i, res in enumerate(results):
            pair = res.get('pair', '?')
            table.rows[i+1].cells[0].text = pair
            table.rows[i+1].cells[1].text = f'{res.get("lev","?")}x'
            table.rows[i+1].cells[2].text = f'{res.get("pos",0)*100:.0f}%'
            table.rows[i+1].cells[3].text = f'{res.get("sl",0)*100:.0f}%'
            table.rows[i+1].cells[4].text = f'${res.get("final",0):,.2f}'
            ret = res.get("return_pct", 0)
            table.rows[i+1].cells[5].text = f'{ret:+.1f}%'
            table.rows[i+1].cells[6].text = f'{res.get("max_dd_pct",0):.1f}%'
            table.rows[i+1].cells[7].text = f'{res.get("win_rate",0):.1f}%'

            if ret > 0:
                table.rows[i+1].cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0,128,0)
            else:
                table.rows[i+1].cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(200,0,0)

        doc.add_paragraph('')

    # Best configs
    best_btc = data.get('best_btc', {})
    best_eth = data.get('best_eth', {})
    if best_btc or best_eth:
        doc.add_heading('二、最优配置', 1)
        if best_btc:
            doc.add_paragraph(f'BTC/USDT: {best_btc.get("lev","?")}x {best_btc.get("pos",0)*100:.0f}%仓 {best_btc.get("sl",0)*100:.0f}%止损 → {best_btc.get("return_pct",0):+.1f}% 回撤{best_btc.get("max_dd_pct",0):.1f}%')
        if best_eth:
            doc.add_paragraph(f'ETH/USDT: {best_eth.get("lev","?")}x {best_eth.get("pos",0)*100:.0f}%仓 {best_eth.get("sl",0)*100:.0f}%止损 → {best_eth.get("return_pct",0):+.1f}% 回撤{best_eth.get("max_dd_pct",0):.1f}%')

    # Risk warning
    doc.add_paragraph('')
    doc.add_heading('三、风险警告', 1)
    p = doc.add_paragraph()
    run = p.add_run('历史回测不代表未来表现。加密货币交易存在重大亏损风险。')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)

    doc.save(str(docx_path))
    return True

# Convert all JSON reports
json_files = sorted(REPORTS_DIR.glob('*.json'))
converted = 0
for jf in json_files:
    dx = REPORTS_DIR / (jf.stem + '.docx')
    if convert_json_to_word(jf, dx):
        converted += 1
        print(f'  {jf.name} -> {dx.name}')

print(f'\nConverted {converted} reports to Word format.')
