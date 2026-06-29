"""Generate Chinese project manual v2"""
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

doc.add_heading('以太 AI Trader — 项目说明书', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('AI 驱动的双信号加密货币永续合约交易系统').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'版本 2.0 | {datetime.now().strftime("%Y-%m-%d")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('LightGBM + FinRL PPO 强化学习 双信号架构')
doc.add_paragraph('MCP 协议支持 AI 操作员远程控制')
doc.add_page_break()

doc.add_heading('目录', 1)
for s in ['1. 项目概述','2. 系统架构','3. AI 决策管线','4. MCP 控制协议','5. 安全规则',
           '6. Web 仪表盘','7. 交易引擎','8. 回测结果','9. 安装部署','10. 使用指南',
           '11. 项目结构','12. 常见问题','13. 风险警告']:
    doc.add_paragraph(s)

doc.add_heading('1. 项目概述', 1)
doc.add_paragraph('以太 AI Trader 是一个 AI 驱动的加密货币永续合约自动交易系统。使用双信号架构（LightGBM 梯度提升 + FinRL PPO 强化学习），在 OKX 交易所自动交易 BTC/USDT 和 ETH/USDT 永续合约。')
doc.add_paragraph('')
doc.add_paragraph('核心特性：')
doc.add_paragraph('双信号 AI：LightGBM（监督学习）+ FinRL PPO（强化学习）提供两路独立交易信号', style='List Bullet')
doc.add_paragraph('MCP 控制：通过 9 个 MCP 工具，Claude AI 可以远程监控和控制系统', style='List Bullet')
doc.add_paragraph('零人工干预：无需选择策略、调参数或看 K 线图', style='List Bullet')
doc.add_paragraph('自进化：每 4 小时自动重训练，模型热替换', style='List Bullet')
doc.add_paragraph('10 条安全规则：硬编码，AI 不可越权', style='List Bullet')
doc.add_paragraph('交易存档：每笔决策和交易本地归档', style='List Bullet')

doc.add_heading('2. 系统架构', 1)
doc.add_paragraph('系统由 4 个主要层级组成：')

doc.add_heading('2.1 信号层（双信号）', 2)
doc.add_paragraph('信号 A — LightGBM（监督学习）：51 个技术指标 → LightGBM 回归器 → 预期收益率 + 置信度。基于 18 个月 OKX 4h K 线训练。', style='List Bullet')
doc.add_paragraph('信号 B — FinRL PPO（强化学习）：OHLCV 数据 → Gymnasium 交易环境 → PPO 智能体 → {做多, 做空, 观望}。通过 stable-baselines3 训练，80,000 步。', style='List Bullet')
doc.add_paragraph('融合：两路信号同时输入决策仲裁器。RL 信号有否决权。LightGBM 提供置信度分数。', style='List Bullet')

doc.add_heading('2.2 决策层', 2)
doc.add_paragraph('市场状态分类器：6 种市场状态（趋势强/弱、震荡窄/宽、高/低波动）', style='List Bullet')
doc.add_paragraph('风险计算器：仓位大小、止损价、止盈价、杠杆倍数', style='List Bullet')
doc.add_paragraph('决策仲裁器：融合信号 + 安全规则 → 最终操作', style='List Bullet')
doc.add_paragraph('EMA-Trend 过滤器：仅在趋势方向交易', style='List Bullet')

doc.add_heading('2.3 控制层（MCP）', 2)
doc.add_paragraph('API 桥接：FastAPI 服务器（端口 8081），提供 REST + MCP 端点', style='List Bullet')
doc.add_paragraph('9 个 MCP 工具：状态查询、决策查看、统计、信号注入、参数覆写、强制训练、健康检查、交易日查询', style='List Bullet')
doc.add_paragraph('Claude 集成：Claude Code 通过 MCP 连接，监控和控制交易', style='List Bullet')

doc.add_heading('2.4 执行层', 2)
doc.add_paragraph('LiveTrader：独立同步 ccxt 交易引擎（不依赖 freqtrade）', style='List Bullet')
doc.add_paragraph('Freqtrade 桥接：AIStrategy 兼容 freqtrade（可选）', style='List Bullet')
doc.add_paragraph('TradeJournal：JSONL 归档每笔决策和交易', style='List Bullet')

doc.add_heading('3. AI 决策管线', 1)

doc.add_heading('3.1 第一层 — 市场状态分类器', 2)
doc.add_paragraph('模型：LightGBM 6 分类器')
doc.add_paragraph('输入：51 个技术指标')
doc.add_paragraph('输出：六种市场状态之一')
doc.add_paragraph('训练：30 天滚动窗口')

doc.add_heading('3.2 第二层A — 方向预测器（LightGBM）', 2)
doc.add_paragraph('模型：LightGBM 回归器')
doc.add_paragraph('输出：预期收益率（未来4h）、置信度（每样本）、预期最大回撤')
doc.add_paragraph('训练：60 天滚动窗口，时序 80/20 分割')

doc.add_heading('3.3 第二层B — RL 信号（FinRL PPO）', 2)
doc.add_paragraph('框架：stable-baselines3 PPO + Gymnasium 交易环境')
doc.add_paragraph('训练：80,000 步，4h OHLCV 数据')
doc.add_paragraph('输出：LONG / SHORT / HOLD（无置信度分数）')
doc.add_paragraph('定位：RL 提供第二独立意见；LightGBM 仍是主信号')

doc.add_heading('3.4 第三层 — 风险计算器', 2)
doc.add_paragraph('计算：最大仓位、止损价（基于 ATR）、止盈价、杠杆倍数')
doc.add_paragraph('市场状态乘数：强趋势 1.0，弱趋势 0.7，震荡市 0.0（禁止），高波动 0.0（禁止）')
doc.add_paragraph('单笔亏损上限：仓位 8%')
doc.add_paragraph('绝对仓位上限：$500')

doc.add_heading('3.5 第四层 — 决策仲裁器', 2)
doc.add_paragraph('融合 LightGBM 信号 + RL 信号 + 市场状态 + 风险计算 → 最终操作')
doc.add_paragraph('10 条安全规则级联应用')
doc.add_paragraph('RL 可对高置信度分歧行使否决权')

doc.add_heading('4. MCP 控制协议', 1)
doc.add_paragraph('Claude Code 可通过 API 桥接（端口 8081）的 9 个 MCP 工具控制交易系统：')
doc.add_paragraph('')
doc.add_paragraph('GET 端点：')
doc.add_paragraph('/api/v1/ai/status — 模型版本、训练状态、自适应参数', style='List Bullet')
doc.add_paragraph('/api/v1/ai/decision — 最新 AI 交易决策', style='List Bullet')
doc.add_paragraph('/api/v1/ai/stats — 优化器交易统计', style='List Bullet')
doc.add_paragraph('/api/v1/ai/training — 训练调度器状态', style='List Bullet')
doc.add_paragraph('/api/v1/ai/health — 健康检查', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('POST 端点（MCP 控制）：')
doc.add_paragraph('/api/v1/ai/signal — 注入外部信号', style='List Bullet')
doc.add_paragraph('/api/v1/ai/override — 覆写参数（置信度阈值、仓位倍率等）', style='List Bullet')
doc.add_paragraph('/api/v1/ai/train — 强制触发模型重训练', style='List Bullet')

doc.add_heading('5. 安全规则（10 条）', 1)
rules_cn = [
    ('高波动市场', '禁止开新仓'),
    ('震荡市场', '禁止交易（实测亏损 -55% 至 -90%）'),
    ('置信度 < 55%', '不动（自适应阈值，连亏时自动收紧）'),
    ('预期回撤 > 5% 权益', '不动'),
    ('已有亏损仓位', '不开同方向新仓'),
    ('极端资金费率', '方向限制'),
    ('连续 3 笔亏损', '停机 12 小时'),
    ('最大仓位', '20% 权益（绝对上限 $500）'),
    ('最大杠杆', '默认 5x（测试用 10x）'),
    ('单笔止损', '仓位 8%'),
]
for i, (rule, desc) in enumerate(rules_cn):
    doc.add_paragraph(f'规则 {i+1} — {rule}：{desc}', style='List Bullet')

doc.add_heading('6. Web 仪表盘', 1)
doc.add_paragraph('React + TypeScript + Tailwind + Recharts。访问 http://localhost:3000')
doc.add_paragraph('')
doc.add_paragraph('顶部卡片：总权益、今日盈亏、持仓数、AI 状态')
doc.add_paragraph('第二行卡片：当前市场状态、仓位倍率、RL 状态、MCP 状态')
doc.add_paragraph('其他面板：权益曲线（7天）、持仓表、交易记录、AI 决策、训练状态、启动检查清单、控制栏')

doc.add_heading('7. 交易引擎', 1)
doc.add_paragraph('LiveTrader（engine/live_trader.py）是主执行引擎，不依赖 freqtrade。')
doc.add_paragraph('')
doc.add_paragraph('事件循环（每 3 分钟）：')
doc.add_paragraph('1. 从 OKX 获取最新 300 根 OHLCV K 线', style='List Bullet')
doc.add_paragraph('2. 计算 51 个技术指标', style='List Bullet')
doc.add_paragraph('3. 运行方向预测器 + 可选 RL 信号', style='List Bullet')
doc.add_paragraph('4. 应用 EMA-Trend 过滤器', style='List Bullet')
doc.add_paragraph('5. 决策仲裁器产出最终操作', style='List Bullet')
doc.add_paragraph('6. 决策存档到 journal/', style='List Bullet')
doc.add_paragraph('7. 实盘模式：在 OKX 下单；模拟模式：仅记录', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('当前配置：杠杆 10x | 仓位 20% | 止损 8% | 最小信号 0.1% | 最低置信度 60% | 间隔 3 分钟')

doc.add_heading('8. 回测结果', 1)
doc.add_paragraph('所有测试使用真实 OKX 4h K 线（3,265 根/币，18 个月）。')
doc.add_paragraph('')
doc.add_paragraph('风控测试（200 笔，EMA-Trend 过滤）：')
doc.add_paragraph('BTC 保守：+25.3%，回撤 5.7%，胜率 65.0%', style='List Bullet')
doc.add_paragraph('BTC 进取：+62.5%，回撤 13.3%，胜率 65.0%', style='List Bullet')
doc.add_paragraph('ETH 保守：+55.0%，回撤 6.2%，胜率 62.0%', style='List Bullet')
doc.add_paragraph('ETH 进取：+167.8%，回撤 14.1%，胜率 62.0%', style='List Bullet')
doc.add_paragraph('全部 6 配置存活，零爆仓', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('RL + LightGBM 双信号测试：ETH/USDT 牛市 +62%、熊市 +302% 提升')
doc.add_paragraph('')
doc.add_paragraph('蒙特卡洛模拟：ETH 破产概率 0.0%（1000 次模拟）')
doc.add_paragraph('')
doc.add_paragraph('前向验证（6 个季度）：BTC 平均 +18.3%/季，ETH 平均 +110.1%/季，全部 12/12 盈利')

doc.add_heading('9. 安装部署', 1)
doc.add_paragraph('')
doc.add_paragraph('环境要求：Python 3.11+、Node.js 20+（仪表盘）、OKX 账号 + API Key、v2rayN 或等效代理')
doc.add_paragraph('')
doc.add_paragraph('快速开始：')
doc.add_paragraph('git clone <仓库地址>')
doc.add_paragraph('cd ethereum-ai-trader')
doc.add_paragraph('pip install -r requirements.txt')
doc.add_paragraph('cp .env.example .env')
doc.add_paragraph('')
doc.add_paragraph('编辑 .env 填入 OKX API 凭证：')
doc.add_paragraph('OKX_API_KEY=你的密钥')
doc.add_paragraph('OKX_API_SECRET=你的密钥')
doc.add_paragraph('OKX_API_PASSPHRASE=你的密码')
doc.add_paragraph('')
doc.add_paragraph('python -m engine.trainer    # 训练模型')
doc.add_paragraph('python -m engine.live_trader  # 启动模拟交易')

doc.add_heading('10. 使用指南', 1)
doc.add_paragraph('')
doc.add_paragraph('交易模式：')
doc.add_paragraph('模拟交易（安全）：python -m engine.live_trader', style='List Bullet')
doc.add_paragraph('实盘交易（真实资金！）：python -m engine.live_trader --live', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('Web 仪表盘：cd web && npm install && npm run dev，打开 http://localhost:3000')
doc.add_paragraph('')
doc.add_paragraph('API 桥接 + MCP 控制：python -m engine.api_bridge --port 8081')
doc.add_paragraph('')
doc.add_paragraph('AI 操作员命令：')
doc.add_paragraph('python -m engine.ai_operator status — 系统状态', style='List Bullet')
doc.add_paragraph('python -m engine.ai_operator trades --last 10 — 最近交易', style='List Bullet')
doc.add_paragraph('python -m engine.ai_operator daily — 今日汇总', style='List Bullet')
doc.add_paragraph('python -m engine.ai_operator check — 异常检测', style='List Bullet')
doc.add_paragraph('python -m engine.ai_operator adjust --confidence 0.65 — 调整风险参数', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('Claude Code 集成：输入 /crypto-trader 启动 AI 交易操作员代理', style='List Bullet')

doc.add_heading('11. 项目结构', 1)
doc.add_paragraph('engine/ — AI 交易引擎（21 个模块）')
doc.add_paragraph('web/ — React 仪表盘（7 个组件）')
doc.add_paragraph('tests/ — 测试套件（20+ 文件）')
doc.add_paragraph('models/ — 训练好的模型文件（.gitignored）')
doc.add_paragraph('journal/ — 交易记录（.gitignored）')
doc.add_paragraph('reports/ — 测试报告（.gitignored）')
doc.add_paragraph('requirements.txt — Python 依赖')
doc.add_paragraph('.env.example — API 密钥模板')
doc.add_paragraph('start.sh / start.bat — 一键启动脚本')

doc.add_heading('12. 常见问题', 1)
faq_cn = [
    ('问：需要 Claude Code 或其他 AI 才能运行吗？', '答：不需要。内置的 LightGBM 和 RL 模型在本地 CPU 运行。Claude Code 可选，用于 MCP 监控/控制。'),
    ('问：最少需要多少资金？', '答：ETH 最低约 $60（0.1 ETH，3x 杠杆）。推荐 $500+。'),
    ('问：回测盈利能保证实盘也赚钱吗？', '答：不能保证。历史表现不代表未来结果。仅供教育目的。'),
    ('问：为什么只做 BTC 和 ETH？', '答：流动性最好，庄家操控概率最低，数据最可靠。'),
    ('问：什么是 RL 双信号？', '答：FinRL PPO 强化学习提供第二路独立交易意见。实测 ETH 牛市 +62% 提升。'),
    ('问：Claude 能控制交易吗？', '答：可以。通过 MCP 协议（9 个工具），Claude 可以注入信号、覆写参数、触发训练、检查状态。'),
]
for q, a in faq_cn:
    doc.add_paragraph(q)
    doc.add_paragraph(a)
    doc.add_paragraph('')

doc.add_heading('13. 风险警告', 1)
p = doc.add_paragraph()
run = p.add_run('警告：')
run.bold = True; run.font.color.rgb = RGBColor(200, 0, 0)
p.add_run('本系统仅供教育目的。加密货币交易存在重大亏损风险。历史表现不代表未来结果。10x 杠杆意味着 10% 的反向波动即可导致全部亏损。永远不要投入无法承受损失的资金。作者不对任何交易结果承担责任。')

doc.add_paragraph('')
doc.add_paragraph(f'文档生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
doc.add_paragraph('项目版本：2.0 | Git 提交：33+ | 测试覆盖：100+ 测试')

out = 'PROJECT_MANUAL_V2_CN.docx'
doc.save(out)
print(f'Saved: {out}')
