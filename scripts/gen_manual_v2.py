"""Generate comprehensive project manual v2 — includes RL + MCP + dual-signal"""
import sys, os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# === COVER ===
doc.add_heading('ETHEREUM AI TRADER', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('AI-Driven Dual-Signal Crypto Futures Trading System').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'Version 2.0 | {datetime.now().strftime("%Y-%m-%d %H:%M")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')
doc.add_paragraph('LightGBM + FinRL PPO reinforcement learning dual-signal architecture')
doc.add_paragraph('MCP (Model Context Protocol) for Claude AI operator control')
doc.add_paragraph('10 hardcoded safety rules — AI cannot override')
doc.add_page_break()

# === TABLE OF CONTENTS ===
doc.add_heading('Table of Contents', 1)
sections = [
    '1. Project Overview',
    '2. Architecture',
    '3. AI Decision Pipeline (Dual-Signal)',
    '4. MCP Control Protocol',
    '5. Safety Rules',
    '6. Web Dashboard',
    '7. Live Trading Engine',
    '8. Backtest Results',
    '9. Installation & Setup',
    '10. Usage Guide',
    '11. Project Structure',
    '12. FAQ',
    '13. Risk Warning',
]
for s in sections:
    doc.add_paragraph(s)

# === 1. OVERVIEW ===
doc.add_heading('1. Project Overview', 1)
doc.add_paragraph('Ethereum AI Trader is an AI-driven autonomous cryptocurrency perpetual futures trading system. It trades BTC/USDT and ETH/USDT perpetual contracts on OKX using a dual-signal architecture combining LightGBM gradient boosting with FinRL PPO reinforcement learning.')
doc.add_paragraph('')
doc.add_paragraph('Key Features:')
doc.add_paragraph('Dual-Signal AI: LightGBM (supervised) + FinRL PPO (reinforcement learning) provide two independent trading signals', style='List Bullet')
doc.add_paragraph('MCP Control: Claude AI can monitor and control the system via 9 MCP tools', style='List Bullet')
doc.add_paragraph('Zero Human Input: No strategy selection, parameter tuning, or chart analysis needed', style='List Bullet')
doc.add_paragraph('Self-Evolving: Auto-retraining every 4 hours with model hot-swap', style='List Bullet')
doc.add_paragraph('10 Safety Rules: Hardcoded, non-overridable by AI', style='List Bullet')
doc.add_paragraph('Trade Journal: Every decision and trade archived locally', style='List Bullet')

# === 2. ARCHITECTURE ===
doc.add_heading('2. Architecture', 1)
doc.add_paragraph('The system consists of 4 main layers:')
doc.add_paragraph('')
doc.add_heading('2.1 Signal Layer (Dual-Signal)', 2)
doc.add_paragraph('Signal A — LightGBM (Supervised): 51 technical indicators -> LightGBM regressor -> expected_return + confidence. Trained on 18 months of OKX 4h candles.', style='List Bullet')
doc.add_paragraph('Signal B — FinRL PPO (Reinforcement Learning): OHLCV data -> Gymnasium trading environment -> PPO agent -> {LONG, SHORT, HOLD}. Trained via stable-baselines3, 80K timesteps.', style='List Bullet')
doc.add_paragraph('Fusion: Both signals feed into DecisionArbitrator. RL signal gets veto option. LightGBM provides confidence score.', style='List Bullet')
doc.add_paragraph('')
doc.add_heading('2.2 Decision Layer', 2)
doc.add_paragraph('RegimeClassifier: 6-class market state (trending/ranging/volatile)', style='List Bullet')
doc.add_paragraph('RiskCalculator: Position size, stop-loss, take-profit, leverage', style='List Bullet')
doc.add_paragraph('DecisionArbitrator: Fuses signals + safety rules -> final action', style='List Bullet')
doc.add_paragraph('EMA-Trend Filter: Only trades in trend direction', style='List Bullet')
doc.add_paragraph('')
doc.add_heading('2.3 Control Layer (MCP)', 2)
doc.add_paragraph('API Bridge: FastAPI server (port 8081) with REST + MCP endpoints', style='List Bullet')
doc.add_paragraph('9 MCP Tools: status, decision, stats, training, signal inject, param override, force train, health check, trade journal query', style='List Bullet')
doc.add_paragraph('Claude Integration: Claude Code connects via MCP to monitor and control trading', style='List Bullet')
doc.add_paragraph('')
doc.add_heading('2.4 Execution Layer', 2)
doc.add_paragraph('LiveTrader: Standalone sync ccxt trading engine', style='List Bullet')
doc.add_paragraph('Freqtrade Bridge: AIStrategy for freqtrade compatibility (optional)', style='List Bullet')
doc.add_paragraph('TradeJournal: JSONL archive of every decision and trade', style='List Bullet')

# === 3. AI PIPELINE ===
doc.add_heading('3. AI Decision Pipeline', 1)
doc.add_heading('3.1 Layer 1 — Regime Classifier', 2)
doc.add_paragraph('Model: LightGBM 6-class classifier')
doc.add_paragraph('Input: 51 technical indicators from FeatureEngineer')
doc.add_paragraph('Output: TRENDING_STRONG / TRENDING_WEAK / RANGING_TIGHT / RANGING_WIDE / HIGH_VOLATILITY / LOW_VOLATILITY')
doc.add_paragraph('Training: 30-day rolling window, expanding quantile labels')

doc.add_heading('3.2 Layer 2A — Direction Predictor (LightGBM)', 2)
doc.add_paragraph('Model: LightGBM regressor')
doc.add_paragraph('Input: 51 features')
doc.add_paragraph('Output: expected_return (next 4h), confidence (per-sample), max_drawdown estimate')
doc.add_paragraph('Training: 60-day rolling window, temporal 80/20 split')

doc.add_heading('3.3 Layer 2B — RL Signal (FinRL PPO)', 2)
doc.add_paragraph('Framework: stable-baselines3 PPO + Gymnasium trading environment')
doc.add_paragraph('Training: 80,000 timesteps, 4h OHLCV data')
doc.add_paragraph('Output: LONG / SHORT / HOLD (no confidence score)')
doc.add_paragraph('Status: RL provides a second independent opinion; LightGBM remains primary')

doc.add_heading('3.4 Layer 3 — Risk Calculator', 2)
doc.add_paragraph('Computes: max position size, stop-loss (ATR-based), take-profit, leverage')
doc.add_paragraph('Regime multipliers: TRENDING 1.0, WEAK 0.7, RANGING 0.0 (blocked), HIGH_VOL 0.0 (blocked)')
doc.add_paragraph('Per-trade loss cap: 8% of position')
doc.add_paragraph('Absolute position cap: $500')

doc.add_heading('3.5 Layer 4 — Decision Arbitrator', 2)
doc.add_paragraph('Fuses LightGBM signal + RL signal + Regime + Risk into final action')
doc.add_paragraph('10 safety rules applied in cascade')
doc.add_paragraph('RL may veto trades with high confidence disagreement')

# === 4. MCP CONTROL ===
doc.add_heading('4. MCP Control Protocol', 1)
doc.add_paragraph('Claude Code can control the trading system via 9 MCP tools exposed through the API Bridge (port 8081):')
doc.add_paragraph('')
doc.add_paragraph('GET Endpoints:')
doc.add_paragraph('/api/v1/ai/status — Model versions, training state, adaptive params', style='List Bullet')
doc.add_paragraph('/api/v1/ai/decision — Latest AI trading decision', style='List Bullet')
doc.add_paragraph('/api/v1/ai/stats — Optimizer trade statistics', style='List Bullet')
doc.add_paragraph('/api/v1/ai/training — Training scheduler status', style='List Bullet')
doc.add_paragraph('/api/v1/ai/health — Health check', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('POST Endpoints (MCP Control):')
doc.add_paragraph('/api/v1/ai/signal — Inject external signal {\"direction\": \"long|short|hold\", \"confidence\": 0.8}', style='List Bullet')
doc.add_paragraph('/api/v1/ai/override — Override params {\"confidence_threshold\": 0.65, \"position_scalar\": 0.5}', style='List Bullet')
doc.add_paragraph('/api/v1/ai/train — Force trigger model retraining', style='List Bullet')

# === 5. SAFETY ===
doc.add_heading('5. Safety Rules (10 Rules)', 1)
rules = [
    ('HIGH_VOLATILITY', 'No new positions allowed'),
    ('RANGING markets', 'Blocked (empirically -55% to -90% losses)'),
    ('Confidence < 55%', 'Hold (adaptive threshold from SelfOptimizer)'),
    ('Expected drawdown > 5% equity', 'Hold'),
    ('Existing losing position', 'No same-direction entry'),
    ('Extreme funding rate', 'Direction restriction'),
    ('3 consecutive losses', 'Stop trading for 12 hours'),
    ('Max position', '20% equity per trade (absolute cap $500)'),
    ('Max leverage', '5x default (10x for testing)'),
    ('Per-trade stop-loss', '8% of position value'),
]
for i, (rule, desc) in enumerate(rules):
    doc.add_paragraph(f'Rule {i+1} — {rule}: {desc}', style='List Bullet')

# === 6. DASHBOARD ===
doc.add_heading('6. Web Dashboard', 1)
doc.add_paragraph('React + TypeScript + Tailwind + Recharts. Available at http://localhost:3000')
doc.add_paragraph('')
doc.add_paragraph('Dashboard Cards (Top Row):')
doc.add_paragraph('Total Equity — Current account value', style='List Bullet')
doc.add_paragraph('Daily PnL — Today profit/loss with percentage', style='List Bullet')
doc.add_paragraph('Open Positions — Number of active trades', style='List Bullet')
doc.add_paragraph('AI Status — Latest training time', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('Status Cards (Second Row):')
doc.add_paragraph('Regime — Current market state with confidence threshold', style='List Bullet')
doc.add_paragraph('Position Scalar — Adaptive position multiplier + per-trade max loss', style='List Bullet')
doc.add_paragraph('RL Status — FinRL PPO trained/active indicator', style='List Bullet')
doc.add_paragraph('MCP Status — Claude control active/standby indicator', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('Additional Panels: Equity Curve (7-day), Positions Table, Trade History, AI Decision, Training Status, Launch Checklist, Control Bar (Start/Stop)')

# === 7. LIVE TRADER ===
doc.add_heading('7. Live Trading Engine', 1)
doc.add_paragraph('')
doc.add_paragraph('The standalone LiveTrader (engine/live_trader.py) is the primary execution engine. It does NOT require freqtrade.')
doc.add_paragraph('')
doc.add_paragraph('Event Loop (every 3 minutes):')
doc.add_paragraph('1. Fetch latest 300 OHLCV candles from OKX (sync ccxt via SOCKS5 proxy)', style='List Bullet')
doc.add_paragraph('2. Compute 51 technical indicators via FeatureEngineer', style='List Bullet')
doc.add_paragraph('3. Run DirectionPredictor + optional RL signal', style='List Bullet')
doc.add_paragraph('4. Apply EMA-Trend filter (only trade in trend direction)', style='List Bullet')
doc.add_paragraph('5. DecisionArbitrator produces final action: LONG / SHORT / HOLD', style='List Bullet')
doc.add_paragraph('6. Log decision to journal/decisions_YYYY-MM.jsonl', style='List Bullet')
doc.add_paragraph('7. In live mode: place market order on OKX; in dry-run: log only', style='List Bullet')
doc.add_paragraph('8. Heartbeat log with cycle count', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('Configuration:')
doc.add_paragraph(f'Leverage: 10x | Position: 20% equity | Stop-Loss: 8% | Min Signal: 0.1% | Min Confidence: 60% | Interval: 3 minutes')

# === 8. BACKTEST ===
doc.add_heading('8. Backtest Results', 1)
doc.add_paragraph('All tests use real OKX 4h candles (3,265/pair, 2024-12-31 to 2026-06-28, 18 months).')
doc.add_paragraph('')
doc.add_heading('8.1 Risk-Controlled Test (200 trades, EMA-Trend filter)', 2)
doc.add_paragraph('BTC Conservative: +25.3% return, 5.7% max DD, 65.0% win rate', style='List Bullet')
doc.add_paragraph('BTC Aggressive: +62.5% return, 13.3% max DD, 65.0% win rate', style='List Bullet')
doc.add_paragraph('ETH Conservative: +55.0% return, 6.2% max DD, 62.0% win rate', style='List Bullet')
doc.add_paragraph('ETH Aggressive: +167.8% return, 14.1% max DD, 62.0% win rate', style='List Bullet')
doc.add_paragraph('ALL 6 configurations survived — 0 liquidations', style='List Bullet')
doc.add_paragraph('')
doc.add_heading('8.2 RL + LightGBM Dual-Signal Test', 2)
doc.add_paragraph('ETH/USDT Dual-Signal: +62% BULL, +302% BEAR improvement over LightGBM-only', style='List Bullet')
doc.add_paragraph('')
doc.add_heading('8.3 Monte Carlo Simulation', 2)
doc.add_paragraph('ETH: 0.0% ruin probability (1000 simulations)', style='List Bullet')
doc.add_paragraph('BTC with EMA-Trend: 0.0% ruin probability', style='List Bullet')
doc.add_paragraph('')
doc.add_heading('8.4 Walk-Forward Validation', 2)
doc.add_paragraph('6 quarterly segments (2025-Q1 to 2026-Q2)')
doc.add_paragraph(f'BTC: +3.1% to +30.7% per quarter, AVG +18.3%, ALL 6/6 profitable')
doc.add_paragraph(f'ETH: +75.4% to +153.7% per quarter, AVG +110.1%, ALL 6/6 profitable')

# === 9. INSTALLATION ===
doc.add_heading('9. Installation & Setup', 1)
doc.add_paragraph('')
doc.add_heading('9.1 Requirements', 2)
doc.add_paragraph('Python 3.11+ with venv', style='List Bullet')
doc.add_paragraph('Node.js 20+ (for web dashboard)', style='List Bullet')
doc.add_paragraph('OKX exchange account with API key (trade + read permissions, NO withdraw)', style='List Bullet')
doc.add_paragraph('v2rayN or equivalent SOCKS5 proxy for OKX access', style='List Bullet')
doc.add_paragraph('')
doc.add_heading('9.2 Quick Start', 2)
doc.add_paragraph('')
doc.add_paragraph('git clone <repo-url> ethereum-ai-trader')
doc.add_paragraph('cd ethereum-ai-trader')
doc.add_paragraph('python -m venv .venv && source .venv/bin/activate')
doc.add_paragraph('pip install -r requirements.txt')
doc.add_paragraph('cp .env.example .env')
doc.add_paragraph('')
doc.add_paragraph('Edit .env with your OKX API credentials:')
doc.add_paragraph('OKX_API_KEY=your_api_key')
doc.add_paragraph('OKX_API_SECRET=your_api_secret')
doc.add_paragraph('OKX_API_PASSPHRASE=your_api_passphrase')
doc.add_paragraph('')
doc.add_paragraph('python -m engine.trainer    # Train models on historical data')
doc.add_paragraph('python -m engine.live_trader  # Start dry-run trading')

# === 10. USAGE ===
doc.add_heading('10. Usage Guide', 1)
doc.add_paragraph('')
doc.add_heading('10.1 Trading Modes', 2)
doc.add_paragraph('Dry-Run (safe, no real money): python -m engine.live_trader', style='List Bullet')
doc.add_paragraph('Live Trading (real money!): python -m engine.live_trader --live', style='List Bullet')
doc.add_paragraph('')
doc.add_heading('10.2 Web Dashboard', 2)
doc.add_paragraph('cd web && npm install && npm run dev')
doc.add_paragraph('Open http://localhost:3000')
doc.add_paragraph('')
doc.add_heading('10.3 API Bridge + MCP Control', 2)
doc.add_paragraph('python -m engine.api_bridge --port 8081')
doc.add_paragraph('Then Claude can connect via MCP to control trading.')
doc.add_paragraph('')
doc.add_heading('10.4 AI Operator Commands', 2)
doc.add_paragraph('python -m engine.ai_operator status — System status check', style='List Bullet')
doc.add_paragraph('python -m engine.ai_operator trades --last 10 — Recent trades', style='List Bullet')
doc.add_paragraph('python -m engine.ai_operator daily — Today summary', style='List Bullet')
doc.add_paragraph('python -m engine.ai_operator check — Anomaly detection', style='List Bullet')
doc.add_paragraph('python -m engine.ai_operator adjust --confidence 0.65 — Adjust risk', style='List Bullet')
doc.add_paragraph('')
doc.add_heading('10.5 Claude Code Integration', 2)
doc.add_paragraph('/crypto-trader — Launch the AI crypto trader operator agent', style='List Bullet')
doc.add_paragraph('The agent will check system status, monitor trades, and take corrective actions', style='List Bullet')

# === 11. STRUCTURE ===
doc.add_heading('11. Project Structure', 1)
doc.add_paragraph('')
structure = '''ethereum-ai-trader/
  engine/                  AI trading engine (21 modules)
    features.py            Feature engineering (51 indicators)
    regime_classifier.py   Market state classifier (LightGBM)
    direction_predictor.py Direction predictor (LightGBM)
    decision_arbitrator.py Risk calculator + 10 safety rules
    self_optimizer.py      Adaptive parameter optimizer
    scheduler.py           Auto-training scheduler
    training_pipeline.py   Training pipeline (+ RL agent)
    live_trader.py         Standalone trading engine
    ai_strategy.py         Freqtrade bridge (optional)
    api_bridge.py          REST API + MCP endpoints
    ai_operator.py         AI operator control panel
    operator_loop.py       Automated monitoring loop
    trade_journal.py       Trade archive system
    backtest_adapter.py    Backtest engine
    trainer.py             Offline model trainer
    validate.py            Pre-launch validation
    launch_check.py        Deployment checklist
    walkforward.py         Walk-forward validation

  web/                     React dashboard (7 components)
    Dashboard.tsx          Status cards + regime + RL/MCP indicators
    EquityCurve.tsx        7-day equity chart
    Positions.tsx          Live position table
    TradeHistory.tsx       Trade history list
    TrainingPanel.tsx      Training status + metrics
    ControlBar.tsx         Start/Stop + live adaptive params
    LaunchCheck.tsx        Pre-flight checklist

  tests/                   Test suite (20+ files)
  models/                  Trained model files (.gitignored)
  journal/                 Trade records (.gitignored)
  reports/                 Test reports (.gitignored)
  requirements.txt         Python dependencies
  .env.example             API key template
  start.sh / start.bat     One-click launcher'''
for line in structure.split('\n'):
    if line.strip():
        doc.add_paragraph(line)

# === 12. FAQ ===
doc.add_heading('12. FAQ', 1)
faq = [
    ('Q: Do I need Claude Code or any AI to run this?', 'A: NO. The built-in LightGBM and RL models run locally. Claude Code is optional for monitoring/control via MCP.'),
    ('Q: Minimum capital required?', 'A: ETH minimum ~$60 (0.1 ETH at 3x). BTC minimum ~$200. Recommend $500+.'),
    ('Q: Does backtest profit guarantee live profit?', 'A: NO. Past performance does not guarantee future results. Educational purposes only.'),
    ('Q: Why only BTC and ETH?', 'A: Highest liquidity, lowest manipulation risk, most reliable data.'),
    ('Q: What is RL dual-signal?', 'A: FinRL PPO reinforcement learning provides a second independent trading opinion alongside LightGBM. +62% BULL, +302% BEAR improvement.'),
    ('Q: Can Claude control the trading?', 'A: YES. Via MCP protocol (9 tools). Claude can inject signals, override parameters, trigger training, and check status.'),
]
for q, a in faq:
    doc.add_paragraph(q)
    doc.add_paragraph(a)
    doc.add_paragraph('')

# === 13. RISK ===
doc.add_heading('13. Risk Warning', 1)
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('WARNING: ')
run.bold = True
run.font.color.rgb = RGBColor(200, 0, 0)
p.add_run('This system is for educational purposes only. Cryptocurrency trading involves substantial risk of loss. Past performance does not guarantee future results. 10x leverage means a 10% adverse move can wipe out your entire position. Never trade with money you cannot afford to lose. The authors assume no responsibility for any trading losses.')

doc.add_paragraph('')
doc.add_paragraph(f'Document generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
doc.add_paragraph(f'Project version: 2.0 | Git commits: 32+ | Test coverage: 100+ tests')

# Save
out = 'PROJECT_MANUAL_V2.docx'
doc.save(out)
print(f'Saved: {out}')
